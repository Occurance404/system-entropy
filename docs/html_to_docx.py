from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_seed_grant_docx():
    doc = Document()
    
    # --- Page Setup ---
    section = doc.sections[0]
    section.page_height = Cm(29.7) # A4 Height
    section.page_width = Cm(21.0)  # A4 Width
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # --- Header with Logo ---
    # Create a table for header layout: Left (Empty) | Right (Logo)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(10)
    header_table.columns[1].width = Cm(7)
    
    # Right cell: Add Logo
    try:
        cell_right = header_table.cell(0, 1)
        paragraph = cell_right.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        run.add_picture('srm_logo.jpg', width=Cm(6)) # Adjust width as needed
    except FileNotFoundError:
        print("Warning: Logo file 'srm_logo.jpg' not found. Skipping image.")

    # --- Title ---
    p = doc.add_paragraph('Annexure -1 Application Form')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    # --- Student Details ---
    p = doc.add_paragraph('Student Details:')
    p.runs[0].bold = True
    
    details_table = doc.add_table(rows=2, cols=2)
    details_table.style = 'Table Grid'
    
    # Row 1
    details_table.cell(0, 0).text = 'Name of Student:'
    details_table.cell(0, 1).text = 'Birat Chapagain'
    
    # Row 2
    details_table.cell(1, 0).text = 'Registration No:'
    details_table.cell(1, 1).text = 'AP24122040024'

    doc.add_paragraph() # Spacer

    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.style = 'Table Grid'
    
    # Row 1
    meta_table.cell(0, 0).text = 'Program:'
    meta_table.cell(0, 1).text = '[PROGRAM]'
    meta_table.cell(0, 2).text = 'School:'
    meta_table.cell(0, 3).text = '[SCHOOL]'
    
    # Row 2
    meta_table.cell(1, 0).text = 'CGPA:'
    meta_table.cell(1, 1).text = '8.44'
    meta_table.cell(1, 2).text = 'Any Backlog:'
    meta_table.cell(1, 3).text = 'No (X)'

    doc.add_paragraph()

    # --- 1. Title of Project ---
    doc.add_paragraph('1. Title of the Project:', style='Heading 2')
    doc.add_paragraph('Entropic Dynamics of Large Language Models Under Non-Stationary Task Constraints')

    # --- 2. Problem Statement ---
    doc.add_paragraph('2. Problem Statement', style='Heading 2')
    doc.add_paragraph('Long-horizon AI coding agents can appear confident while silently drifting from the user intent and task constraints. Standard pass/fail metrics on software engineering tasks do not expose these failures. We need a stability-focused evaluation framework and a context-management strategy that reduces semantic collapse over long interactions.')

    # --- 3. Objectives ---
    doc.add_paragraph('3. Objectives', style='Heading 2')
    obj_p = doc.add_paragraph()
    obj_p.style = 'List Bullet'
    obj_p.add_run('Define a Semantic Collapse Ratio (SCR) metric with probe-based evaluation.\n')
    obj_p.add_run('Build a stability benchmark suite with controlled long-horizon scenarios.\n')
    obj_p.add_run('Develop a toolized context-management agent with structured memory.\n')
    obj_p.add_run('Compare append-only vs context-managed agents on success and stability.\n')
    obj_p.add_run('Prepare a Scopus/WoS-targeted publication with SRM AP acknowledgment.')

    # --- 4. Methodology ---
    doc.add_paragraph('4. Methodology', style='Heading 2')
    meth_p = doc.add_paragraph()
    meth_p.style = 'List Bullet'
    meth_p.add_run('Design and validate SWE scenarios and probes to track semantic drift.\n')
    meth_p.add_run('Run controlled experiments with fixed model, token budgets, and logging.\n')
    meth_p.add_run('Implement a context tool with fixed intent, long-term summary, and working memory.\n')
    meth_p.add_run('Generate stability heatmaps, SCR curves, and pass-rate comparisons.')

    # --- 5. Outcome ---
    doc.add_paragraph('5. Outcome', style='Heading 2')
    out_p = doc.add_paragraph()
    out_p.style = 'List Bullet'
    out_p.add_run('One publication in Scopus/WoS indexed journal or conference proceeding.\n')
    out_p.add_run('Benchmark artifacts and analysis report with SRM AP acknowledgment.')

    doc.add_paragraph()

    # --- Budget Section ---
    doc.add_paragraph('Budget in detail:', style='Heading 2')
    
    # Summary Budget Table
    budget_table = doc.add_table(rows=6, cols=2)
    budget_table.style = 'Table Grid'
    
    rows = [
        ('Field data collection', 'N/A'),
        ('Laboratory chemicals and consumables', 'N/A'),
        ('Characterization/User Charge', 'N/A'),
        ('Charges towards secondary data', 'N/A'),
        ('AI Model API Access Credits\n(Access to commercial AI systems like GPT-5, Claude, Gemini, and DeepSeek via OpenRouter platform for running benchmark experiments)', '₹40,000'),
        ('Total', '₹40,000')
    ]

    for i, (item, cost) in enumerate(rows):
        budget_table.cell(i, 0).text = item
        budget_table.cell(i, 1).text = cost
    
    # Bold the Total and API Credits row
    budget_table.cell(4, 0).paragraphs[0].runs[0].bold = True
    budget_table.cell(4, 1).paragraphs[0].runs[0].bold = True
    budget_table.cell(5, 0).paragraphs[0].runs[0].bold = True
    budget_table.cell(5, 1).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # --- Budget Justification ---
    doc.add_paragraph('Budget Justification', style='Heading 2')
    doc.add_paragraph('Why do we need AI Model API Credits?')
    doc.add_paragraph('This research project evaluates how well AI coding assistants maintain reliability over long, complex tasks. To produce meaningful, publishable results, we must compare multiple state-of-the-art AI systems (such as OpenAI\'s GPT-5, Anthropic\'s Claude, Google\'s Gemini, and DeepSeek) under identical experimental conditions.')
    doc.add_paragraph('Unlike traditional software, commercial AI models are accessed through pay-per-use APIs. Each time we run an experiment, the AI reads our task instructions, file contents, and error messages (input) and generates code, reasoning, and tool calls (output). Both input and output are measured in "tokens" (roughly, chunks of text), and providers charge based on token usage.')

    # --- Detailed Credit Utilization Plan ---
    doc.add_paragraph('Detailed Credit Utilization Plan', style='Heading 2')
    doc.add_paragraph('Each benchmark experiment involves a multi-step coding task where the AI agent reads files, writes code, executes scripts, and iterates on errors. Our experiments are designed to run for 50-100 agent steps, with each step consuming context (accumulated history) plus new output. Additionally, our stability measurement technique requires running 6 parallel "branching probes" at regular intervals to compute the Semantic Collapse Ratio (SCR).')
    doc.add_paragraph('Based on our pilot experiments, each full benchmark run consumes approximately 400,000-500,000 tokens (including input context, output generation, and probe overhead).')

    # Detailed Table
    detail_table = doc.add_table(rows=13, cols=5)
    detail_table.style = 'Table Grid'
    
    headers = ['AI Model', 'Provider', 'Cost per Experiment', 'Experiments', 'Subtotal']
    for col, text in enumerate(headers):
        cell = detail_table.cell(0, col)
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True

    data_rows = [
        ['GPT-5.1', 'OpenAI', '₹250', '15', '₹3,750'],
        ['GPT-5 Mini', 'OpenAI', '₹50', '15', '₹750'],
        ['Claude Sonnet 4.5', 'Anthropic', '₹400', '15', '₹6,000'],
        ['Claude Opus 4.5', 'Anthropic', '₹650', '15', '₹9,750'],
        ['Gemini 3 Pro', 'Google', '₹300', '15', '₹4,500'],
        ['Gemini 2.5 Pro', 'Google', '₹250', '15', '₹3,750'],
        ['DeepSeek-Chat', 'DeepSeek', '₹30', '15', '₹450'],
        ['DeepSeek-Reasoner', 'DeepSeek', '₹60', '15', '₹900']
    ]

    for i, row in enumerate(data_rows, start=1):
        for j, val in enumerate(row):
            detail_table.cell(i, j).text = val

    # Totals
    # Subtotal
    detail_table.cell(9, 3).text = 'Subtotal (Raw API Costs)'
    detail_table.cell(9, 3).paragraphs[0].runs[0].bold = True
    detail_table.cell(9, 4).text = '₹29,850'
    
    # Platform Fee
    detail_table.cell(10, 3).text = 'OpenRouter Platform Fee (5.5%)'
    detail_table.cell(10, 4).text = '₹1,650'
    
    # Failed Runs
    detail_table.cell(11, 3).text = 'Failed Runs & Retries (~20%)'
    detail_table.cell(11, 4).text = '₹6,000'
    
    # Pilot Testing
    detail_table.cell(12, 3).text = 'Pilot Testing & Debugging'
    detail_table.cell(12, 4).text = '₹2,500'

    # Grand Total (Create a new row manually or use text below)
    # Merging cells vertically logic is complex in python-docx simple script, 
    # so we'll add a separate paragraph for the TOTAL.
    
    doc.add_paragraph()
    total_p = doc.add_paragraph('Total Requested: ₹40,000')
    total_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_p.runs[0].bold = True
    total_p.runs[0].font.size = Pt(12)

    doc.add_paragraph('Note: All API access will be routed through OpenRouter, a unified platform that provides a single billing interface for multiple AI providers. This simplifies expense tracking and ensures transparent credit utilization.', style='Quote')

    # --- Financial Support ---
    doc.add_paragraph()
    doc.add_paragraph('Availed Any Financial Support from SRM University AP: No')
    doc.add_paragraph('Provide Details (if yes): [N/A]')

    doc.add_paragraph()
    
    # --- Signature ---
    sig = doc.add_paragraph('Birat C.')
    sig.runs[0].font.name = 'Brush Script MT' # Try a script font, or generic fallback
    sig.runs[0].font.color.rgb = None # Default color
    sig.runs[0].font.size = Pt(20)
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph('__________________________')
    doc.add_paragraph('Signature of Student')

    doc.add_page_break()

    # --- Proposal Summary Attachment ---
    doc.add_paragraph('ATTACHMENT: PROPOSAL SUMMARY', style='Title')
    
    # Abstract
    doc.add_paragraph('Abstract', style='Heading 2')
    doc.add_paragraph('This project introduces a stability-first evaluation for long-horizon software engineering agents. We define a Semantic Collapse Ratio (SCR) metric using probe-based checkpoints to detect when agents drift from task intent while remaining confident. We will build a benchmark suite with controlled long-horizon scenarios.') 
    doc.add_paragraph('We will compare baseline append-only context agents against a structured context-management agent that maintains fixed intent, long-term summaries, and short-term working memory. The outcome will be a reproducible evaluation framework, stability visualizations, and a research paper suitable for Scopus/WoS venues.')

    # Work Plan
    doc.add_paragraph('Work Plan', style='Heading 2')
    wp_p = doc.add_paragraph()
    wp_p.style = 'List Bullet'
    wp_p.add_run('Weeks 1-2: Finalize scenarios, probes, and logging pipeline.\n')
    wp_p.add_run('Weeks 3-4: Run baseline experiments and compute SCR metrics.\n')
    wp_p.add_run('Weeks 5-6: Implement context tool and run comparative experiments.\n')
    wp_p.add_run('Weeks 7-8: Analyze results, prepare paper draft and visuals.')

    # Future Impact
    doc.add_paragraph('Future Impact & Significance', style='Heading 2')
    doc.add_paragraph('If this research succeeds, it will have significant implications for the rapidly growing field of AI-assisted software development:')
    
    fi_p = doc.add_paragraph()
    fi_p.style = 'List Bullet'
    fi_p.add_run('Safer AI Coding Assistants: SCR metric can be integrated into tools like GitHub Copilot to warn developers when AI is unreliable.\n').bold = False
    fi_p.add_run('Industry Adoption: A validated stability benchmark allows enterprises to evaluate AI systems before deployment.\n')
    fi_p.add_run('Foundation for Future Research: The framework can be extended to study different failure modes.\n')
    fi_p.add_run('Improved Context Management: Provides a design pattern for reliable long-running agents.\n')
    fi_p.add_run('Academic Contribution: Addresses the gap in AI stability evaluation literature.')

    doc.add_page_break()

    # --- Recommendations Page ---
    # Header with Logo again
    rec_header_table = doc.add_table(rows=1, cols=2)
    rec_header_table.autofit = False
    rec_header_table.columns[0].width = Cm(10)
    rec_header_table.columns[1].width = Cm(7)
    
    # Right cell: Add Logo
    try:
        cell_right = rec_header_table.cell(0, 1)
        paragraph = cell_right.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        run.add_picture('srm_logo.jpg', width=Cm(6))
    except FileNotFoundError:
        pass

    doc.add_paragraph('Research Supervisor/Faculty Mentor’s Recommendations:', style='Heading 2')
    doc.add_paragraph('[Comments Placeholder]')
    doc.add_paragraph('\n\n')
    doc.add_paragraph('Signature: __________________________    Date: ________________')

    doc.add_paragraph()
    doc.add_paragraph('Recommended by HOD:', style='Heading 2')
    doc.add_paragraph('Signature: __________________________    Date: ________________')

    doc.add_paragraph()
    doc.add_paragraph('Approved by DEAN (Research):', style='Heading 2')
    doc.add_paragraph('Signature: __________________________    Date: ________________')

    doc.save('seed_grant_application.docx')

if __name__ == '__main__':
    create_seed_grant_docx()
